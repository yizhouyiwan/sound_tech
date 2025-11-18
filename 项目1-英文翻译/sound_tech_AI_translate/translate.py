from fastapi import FastAPI, HTTPException
from pydantic import BaseModel
from deepseek import DeepSeekAPI
import os
import json
from config import *
import uvicorn
from docx import Document
from docx.shared import Inches
import uuid
from fastapi.responses import FileResponse
import os

app = FastAPI()

# 创建下载目录
DOWNLOAD_DIR = "downloads"
if not os.path.exists(DOWNLOAD_DIR):
    os.makedirs(DOWNLOAD_DIR)

# 使用正确的类名
# client = DeepSeekAPI(api_key=os.getenv("DEEPSEEK_API_KEY"))
client = DeepSeekAPI(api_key=api_key)

class TranslationRequest(BaseModel):
    text: str
    output_format: str = "json"  # "json" or "word"
    include_vocabulary: bool = True

def create_word_document(original_text: str, translation: str, vocabulary: list, filename: str):
    """
    创建Word文档
    """
    doc = Document()
    
    # 标题
    title = doc.add_heading('智能翻译结果', 0)
    title.alignment = 1  # 居中
    
    # 添加分割线
    doc.add_paragraph("=" * 50)
    
    # 原文部分
    doc.add_heading('英文原文', level=1)
    doc.add_paragraph(original_text)
    
    # 翻译部分
    doc.add_heading('中文翻译', level=1)
    doc.add_paragraph(translation)
    
    # 专业词汇表
    if vocabulary:
        doc.add_heading('专业词汇表', level=1)
        
        # 创建表格：英文 | 中文 | 解释
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        
        # 表头
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '英文词汇'
        hdr_cells[1].text = '中文翻译'
        hdr_cells[2].text = '解释'
        
        # 添加词汇行
        for vocab in vocabulary:
            if isinstance(vocab, dict):
                row_cells = table.add_row().cells
                row_cells[0].text = vocab.get('english', '')
                row_cells[1].text = vocab.get('chinese', '')
                row_cells[2].text = vocab.get('explanation', '')
            else:
                row_cells = table.add_row().cells
                row_cells[0].text = str(vocab)
                row_cells[1].text = ""
                row_cells[2].text = ""
    
    # 保存文档
    file_path = os.path.join(DOWNLOAD_DIR, filename)
    doc.save(file_path)
    return file_path

@app.post(translate_url)
async def translate_text(request: TranslationRequest):
    try:
        prompt = f"""
        请将以下英文文本翻译成中文，并提取重要的专业词汇：

        英文原文：{request.text}

        请严格按照以下JSON格式返回：
        {{
            "translation": "中文翻译内容",
            "vocabulary": [
                {{
                    "english": "专业词汇1",
                    "chinese": "中文翻译1", 
                    "explanation": "详细解释1"
                }},
                {{
                    "english": "专业词汇2",
                    "chinese": "中文翻译2",
                    "explanation": "详细解释2"
                }}
            ]
        }}
        """

        response = client.chat_completion(
            model="deepseek-chat",
            messages=[{"role": "user", "content": prompt}]
        )
        
        try:
            # 解析AI响应
            result = json.loads(response)
            translation = result.get("translation", "")
            vocabulary = result.get("vocabulary", [])
            
            # 生成响应
            response_data = {
                "success": True,
                "translation": translation,
                "vocabulary": vocabulary
            }
            
            # 如果需要Word文档
            if request.output_format == "word":
                filename = f"translation_{uuid.uuid4().hex[:8]}.docx"
                file_path = create_word_document(
                    original_text=request.text,
                    translation=translation,
                    vocabulary=vocabulary,
                    filename=filename
                )
                response_data["word_document_url"] = f"/downloads/{filename}"
            
            return response_data
            
        except Exception as e:
            return {
                "success": False,
                "error": f"解析AI响应失败: {str(e)}",
                "raw_response": response
            }
            
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/downloads/{filename}")
async def download_file(filename: str):
    """
    下载Word文档文件
    """
    file_path = os.path.join(DOWNLOAD_DIR, filename)
    if os.path.exists(file_path):
        return FileResponse(
            path=file_path,
            filename=f"translation_{filename}",
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    else:
        raise HTTPException(status_code=404, detail="文件未找到")

if __name__ == "__main__":
    uvicorn.run(app, host=host, port=port)